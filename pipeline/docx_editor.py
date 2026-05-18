"""Edit a copy of base_resume.docx with tailored content from the LLM.

The strategy: find sections by matching company/project names in paragraph
text, then replace the 3 bullet paragraphs that follow.  We never delete or
add paragraphs — only swap the text content of existing runs so all formatting
(fonts, sizes, bold, spacing, margins) is preserved exactly.

On first use (or when base_resume.docx changes), an auto-calibration step
ensures the resume has breathing room on one page.  Adjustments are applied
directly to base_resume.docx and recorded in a .resume_calibrated flag file.
"""

import hashlib
import json
import logging
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

logger = logging.getLogger(__name__)

RESUME_DIR = Path(__file__).resolve().parent.parent / "resume"
BASE_DOCX = RESUME_DIR / "base_resume.docx"
_CALIBRATION_FLAG = RESUME_DIR / ".resume_calibrated"

# Heading styles are excluded from body-text font reduction.
_HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Title"}

# Margin floor — calibration will reduce margins down to this, never below.
_MARGIN_FLOOR = Inches(0.75)

# Font reduction step (applied once per calibration attempt).
_FONT_STEP = Pt(0.5)


# ---------------------------------------------------------------------------
# Auto-calibration
# ---------------------------------------------------------------------------

def _file_hash(path: Path) -> str:
    """SHA-256 hex digest of a file's contents."""
    h = hashlib.sha256()
    h.update(path.read_bytes())
    return h.hexdigest()


def _needs_calibration() -> bool:
    """Return True if base_resume.docx has never been calibrated or has changed."""
    if not BASE_DOCX.exists():
        return False
    if not _CALIBRATION_FLAG.exists():
        return True
    try:
        flag = json.loads(_CALIBRATION_FLAG.read_text(encoding="utf-8"))
        return flag.get("hash") != _file_hash(BASE_DOCX)
    except (json.JSONDecodeError, OSError):
        return True


def _write_calibration_flag(adjustments: list[str]) -> None:
    """Write the flag file recording the docx hash and what was done."""
    flag = {
        "hash": _file_hash(BASE_DOCX),
        "adjustments": adjustments,
    }
    _CALIBRATION_FLAG.write_text(json.dumps(flag, indent=2), encoding="utf-8")


def _get_page_count(docx_path: Path) -> int | None:
    """Convert docx to PDF via LibreOffice and return page count, or None on error."""
    try:
        import fitz  # pymupdf
    except ImportError:
        return None

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice is None:
        win_path = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")
        if win_path.exists():
            soffice = str(win_path)
        else:
            return None

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf",
                 "--outdir", tmpdir, str(docx_path)],
                capture_output=True, timeout=30,
            )
            if result.returncode != 0:
                return None
            pdf_path = Path(tmpdir) / f"{docx_path.stem}.pdf"
            if not pdf_path.exists():
                return None
            with fitz.open(str(pdf_path)) as pdf_doc:
                return len(pdf_doc)
    except Exception:
        logger.debug("Page count check failed", exc_info=True)
        return None


def _reduce_margins(doc: Document) -> list[str]:
    """Reduce any section margin that exceeds _MARGIN_FLOOR down to _MARGIN_FLOOR.

    Returns a list of human-readable descriptions of what changed.
    """
    changes: list[str] = []
    for i, section in enumerate(doc.sections):
        prefix = f"section {i}: " if len(doc.sections) > 1 else ""
        for attr, label in [
            ("top_margin", "top"),
            ("bottom_margin", "bottom"),
            ("left_margin", "left"),
            ("right_margin", "right"),
        ]:
            current = getattr(section, attr)
            if current is not None and current > _MARGIN_FLOOR:
                old_in = current / 914400
                setattr(section, attr, _MARGIN_FLOOR)
                changes.append(f"{prefix}{label} margin {old_in:.2f}in -> 0.75in")
    return changes


def _reduce_body_font(doc: Document) -> list[str]:
    """Reduce font size by _FONT_STEP on all non-heading runs and styles.

    Returns a list of human-readable descriptions of what changed.
    """
    changes: list[str] = []
    adjusted_styles: set[str] = set()

    # 1. Reduce style-level font sizes (affects paragraphs inheriting the style)
    for style in doc.styles:
        if style.name in _HEADING_STYLES:
            continue
        if not hasattr(style, "font"):
            continue
        if style.font.size is not None:
            old = style.font.size
            style.font.size = old - _FONT_STEP
            adjusted_styles.add(style.name)
            changes.append(
                f"style '{style.name}' font {old / 12700:.1f}pt -> "
                f"{style.font.size / 12700:.1f}pt"
            )

    # 2. Reduce run-level overrides on non-heading paragraphs
    run_count = 0
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name in _HEADING_STYLES:
            continue
        for run in para.runs:
            if run.font.size is not None:
                run.font.size = run.font.size - _FONT_STEP
                run_count += 1

    if run_count:
        changes.append(f"reduced font on {run_count} run-level overrides by {_FONT_STEP / 12700:.1f}pt")

    return changes


def calibrate_base_resume() -> None:
    """Ensure base_resume.docx fits on one page with room for small additions.

    Runs only when .resume_calibrated is absent or its stored hash differs
    from the current base_resume.docx.  Adjustments are applied in order:
      1. Reduce any margin > 0.75 in down to 0.75 in
      2. Reduce body font size by 0.5 pt

    After each step the page count is re-checked; calibration stops as soon
    as the document fits on one page.
    """
    if not _needs_calibration():
        return

    logger.info("Calibrating base_resume.docx for one-page fit...")

    pages = _get_page_count(BASE_DOCX)
    if pages is None:
        logger.warning("Cannot determine page count (LibreOffice or PyMuPDF missing) — skipping calibration")
        return
    if pages == 1:
        logger.info("Base resume is already 1 page — no calibration needed")
        _write_calibration_flag(["none — already 1 page"])
        return

    logger.warning("Base resume renders as %d pages — starting calibration", pages)
    all_adjustments: list[str] = []

    # --- Step 1: margins ---
    doc = Document(str(BASE_DOCX))
    margin_changes = _reduce_margins(doc)
    if margin_changes:
        doc.save(str(BASE_DOCX))
        for c in margin_changes:
            logger.info("Calibration margin: %s", c)
        all_adjustments.extend(margin_changes)

        pages = _get_page_count(BASE_DOCX)
        if pages is not None and pages <= 1:
            logger.info("Calibration complete after margin reduction (%d page)", pages)
            _write_calibration_flag(all_adjustments)
            return
        if pages is not None:
            logger.info("Still %d pages after margin reduction — trying font reduction", pages)
    else:
        logger.info("All margins already <= 0.75in — skipping margin step")

    # --- Step 2: font reduction ---
    doc = Document(str(BASE_DOCX))
    font_changes = _reduce_body_font(doc)
    if font_changes:
        doc.save(str(BASE_DOCX))
        for c in font_changes:
            logger.info("Calibration font: %s", c)
        all_adjustments.extend(font_changes)

        pages = _get_page_count(BASE_DOCX)
        if pages is not None and pages <= 1:
            logger.info("Calibration complete after font reduction (%d page)", pages)
        elif pages is not None:
            logger.warning(
                "Resume still %d pages after all calibration steps — "
                "manual editing may be required", pages,
            )
    else:
        logger.info("No font sizes to reduce — skipping font step")
        if not all_adjustments:
            logger.warning(
                "No calibration adjustments possible but resume is %s pages — "
                "manual editing may be required",
                pages if pages is not None else "unknown",
            )

    _write_calibration_flag(all_adjustments or ["attempted — no effective adjustments"])

# Skills category labels exactly as they appear in the docx
_SKILLS_CATEGORIES = [
    "Languages:",
    "Frameworks & APIs:",
    "Cloud & Infrastructure:",
    "Databases:",
    "Tools & Practices:",
]


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """Replace the visible text of a paragraph while preserving formatting.

    Puts all new text into the first run and clears the rest.  This keeps the
    paragraph style, run font/size/bold from the first run, and avoids
    creating or deleting XML elements.
    """
    runs = paragraph.runs
    if not runs:
        return

    # Put full text in first run, blank out the rest
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


def _find_paragraph_index(paragraphs, search_text: str) -> int | None:
    """Find the index of the first paragraph whose text contains search_text."""
    search_lower = search_text.lower()
    for i, p in enumerate(paragraphs):
        if search_lower in p.text.lower():
            return i
    return None


def _get_original_bullets(paragraphs, anchor_index: int, count: int = 3) -> list[str]:
    """Read the original bullet texts after an anchor paragraph."""
    bullets = []
    idx = anchor_index + 1
    while idx < len(paragraphs) and len(bullets) < count:
        p = paragraphs[idx]
        style_name = p.style.name if p.style else ""
        if style_name == "List Paragraph":
            bullets.append(p.text)
        elif bullets:
            break
        idx += 1
    return bullets


def _replace_bullets_after(paragraphs, anchor_index: int, new_bullets: list[str]) -> bool:
    """Replace the next 3 'List Paragraph' bullets after anchor_index.

    Returns True if all 3 bullets were found and replaced.
    """
    replaced = 0
    idx = anchor_index + 1

    while idx < len(paragraphs) and replaced < len(new_bullets):
        p = paragraphs[idx]
        style_name = p.style.name if p.style else ""

        if style_name == "List Paragraph":
            _replace_paragraph_text(p, new_bullets[replaced])
            replaced += 1
        elif replaced > 0:
            # We started finding bullets but hit a non-bullet paragraph
            break

        idx += 1

    return replaced == len(new_bullets)


def _get_skills_text_and_run(paragraph) -> tuple[int, str] | None:
    """Find the first non-bold run containing skills text in a skills paragraph.

    Returns (run_index, current_skills_text) where current_skills_text is the
    concatenation of all runs from run_index onward.  Returns None if the
    paragraph has no colon or no skills runs.
    """
    runs = paragraph.runs
    if not runs:
        return None

    # Walk through runs tracking character position to find where skills start.
    full_text = paragraph.text
    colon_idx = full_text.find(":")
    if colon_idx == -1:
        return None

    # Skills text starts after ": " (colon plus any whitespace)
    skills_start_char = colon_idx + 1
    while skills_start_char < len(full_text) and full_text[skills_start_char] == " ":
        skills_start_char += 1

    # Map that character offset to a run index
    char_pos = 0
    for i, run in enumerate(runs):
        run_end = char_pos + len(run.text)
        if run_end > skills_start_char:
            skills_text = "".join(r.text for r in runs[i:])
            return (i, skills_text)
        char_pos = run_end

    return None


def _write_skills_text(paragraph, first_skills_run: int, new_skills_text: str) -> None:
    """Write new skills text into the skills runs of a paragraph.

    Puts all text into the first skills run (with bold=False) and blanks out
    every subsequent run.  Label runs before first_skills_run are untouched.
    """
    runs = paragraph.runs
    runs[first_skills_run].text = new_skills_text
    runs[first_skills_run].bold = False
    for r in runs[first_skills_run + 1:]:
        r.text = ""


def _append_skills(paragraphs, additions: list[dict]) -> None:
    """Append new skills to the appropriate category lines in the Skills section.

    Each addition is {"category": "Frameworks & APIs", "skill": "GraphQL"}.
    Only appends — never replaces existing skills.  Preserves bold on the
    category label and keeps skills text non-bold.
    """
    if not additions:
        return

    # Group additions by category
    by_category: dict[str, list[str]] = {}
    for item in additions:
        cat = item.get("category", "")
        skill = item.get("skill", "")
        if cat and skill:
            by_category.setdefault(cat, []).append(skill)

    for p in paragraphs:
        text = p.text.strip()
        for cat_label in _SKILLS_CATEGORIES:
            if not text.startswith(cat_label.replace(":", "")):
                continue

            cat_key = cat_label.rstrip(":")
            new_skills = by_category.get(cat_key, [])
            if not new_skills:
                continue

            result = _get_skills_text_and_run(p)
            if result is None:
                continue
            run_idx, current_skills = result

            # Filter out skills already present (case-insensitive)
            existing_lower = current_skills.lower()
            to_add = [s for s in new_skills if s.lower() not in existing_lower]
            if not to_add:
                logger.debug("All skills already present in %s", cat_key)
                continue

            suffix = ", ".join(to_add)
            new_text = current_skills.rstrip() + ", " + suffix
            _write_skills_text(p, run_idx, new_text)
            logger.info("Appended to %s: %s", cat_key, suffix)
            break


def _would_overflow_one_page(docx_path: Path) -> bool:
    """Check if the docx would render to more than one page.

    Uses _get_page_count (LibreOffice + PyMuPDF).  Returns False on any
    error (fail-open so the pipeline continues).
    """
    pages = _get_page_count(docx_path)
    if pages is None:
        return False
    if pages > 1:
        logger.warning("Resume overflows to %d pages", pages)
        return True
    return False


def _swap_skills(paragraphs, swaps: list[dict]) -> None:
    """Swap low-signal skills with JD-relevant skills in the Skills section.

    Each swap is {"category": "Frameworks & APIs", "remove": "SSE", "add": "REST APIs"}.
    Only performs the swap if the skill to remove is actually found in the line.
    Preserves bold on the category label and keeps skills text non-bold.
    """
    if not swaps:
        return

    # Group swaps by category
    by_category: dict[str, list[dict]] = {}
    for item in swaps:
        cat = item.get("category", "")
        remove = item.get("remove", "")
        add = item.get("add", "")
        if cat and remove and add:
            by_category.setdefault(cat, []).append({"remove": remove, "add": add})

    for p in paragraphs:
        text = p.text.strip()
        for cat_label in _SKILLS_CATEGORIES:
            if not text.startswith(cat_label.replace(":", "")):
                continue

            cat_key = cat_label.rstrip(":")
            swap_list = by_category.get(cat_key, [])
            if not swap_list:
                continue

            result = _get_skills_text_and_run(p)
            if result is None:
                continue
            run_idx, current_skills = result

            new_skills = current_skills
            for swap in swap_list:
                remove_skill = swap["remove"]
                add_skill = swap["add"]

                if add_skill.lower() in new_skills.lower():
                    logger.debug("Skill %r already present in %s, skipping swap", add_skill, cat_key)
                    continue

                if remove_skill in new_skills:
                    new_skills = new_skills.replace(remove_skill, add_skill, 1)
                    logger.info("Swapped in %s: %r -> %r", cat_key, remove_skill, add_skill)
                else:
                    pattern = re.compile(re.escape(remove_skill), re.IGNORECASE)
                    if pattern.search(new_skills):
                        new_skills = pattern.sub(add_skill, new_skills, count=1)
                        logger.info("Swapped in %s: %r -> %r (case-insensitive)", cat_key, remove_skill, add_skill)
                    else:
                        logger.warning("Could not find %r in %s to swap", remove_skill, cat_key)

            if new_skills != current_skills:
                _write_skills_text(p, run_idx, new_skills)
            break


def edit_docx(tailoring_data: dict, company_slug: str, job_id: str) -> Path:
    """Create a tailored .docx from base_resume.docx.

    On the first call (or after the base resume changes), runs
    calibrate_base_resume() to ensure the template fits on one page.

    Args:
        tailoring_data: Parsed JSON from the LLM with experience, projects,
                        skills_additions, and changes keys.
        company_slug:   Company identifier for the filename.
        job_id:         Job identifier for the filename.

    Returns:
        Path to the edited .docx file (caller converts to PDF then deletes).

    Raises:
        FileNotFoundError: If base_resume.docx is missing.
        Exception: Any python-docx or I/O error — caller should fall back.
    """
    if not BASE_DOCX.exists():
        raise FileNotFoundError(f"Base resume not found: {BASE_DOCX}")

    # Auto-calibrate on first use or when the base file changes
    calibrate_base_resume()

    output_dir = Path(__file__).resolve().parent.parent / "outputs"
    output_dir.mkdir(parents=True, exist_ok=True)
    temp_docx = output_dir / f"{company_slug}_{job_id}_tailored.docx"

    # Work on a copy — never touch the original
    shutil.copy2(BASE_DOCX, temp_docx)

    doc = Document(str(temp_docx))
    paragraphs = doc.paragraphs

    # --- Experience bullets ---
    for exp in tailoring_data.get("experience", []):
        company = exp.get("company", "")
        bullets = exp.get("bullets", [])
        if not company or len(bullets) != 3:
            logger.warning("Skipping experience entry: company=%r, bullets=%d", company, len(bullets))
            continue

        idx = _find_paragraph_index(paragraphs, company)
        if idx is None:
            logger.warning("Could not find company paragraph: %r", company)
            continue

        # Length guard: if any new bullet is longer than original, revert to originals
        originals = _get_original_bullets(paragraphs, idx)
        if len(originals) == 3:
            overflow = False
            for orig, new in zip(originals, bullets):
                if len(new.split()) > len(orig.split()):
                    logger.warning(
                        "Bullet too long for %s (%d > %d words), keeping originals",
                        company, len(new.split()), len(orig.split()),
                    )
                    overflow = True
                    break
            if overflow:
                continue

        if not _replace_bullets_after(paragraphs, idx, bullets):
            logger.warning("Could not replace all 3 bullets for %r", company)

    # --- Project bullets ---
    for proj in tailoring_data.get("projects", []):
        name = proj.get("name", "")
        bullets = proj.get("bullets", [])
        if not name or len(bullets) != 3:
            logger.warning("Skipping project entry: name=%r, bullets=%d", name, len(bullets))
            continue

        idx = _find_paragraph_index(paragraphs, name)
        if idx is None:
            logger.warning("Could not find project paragraph: %r", name)
            continue

        # Length guard
        originals = _get_original_bullets(paragraphs, idx)
        if len(originals) == 3:
            overflow = False
            for orig, new in zip(originals, bullets):
                if len(new.split()) > len(orig.split()):
                    logger.warning(
                        "Bullet too long for %s (%d > %d words), keeping originals",
                        name, len(new.split()), len(orig.split()),
                    )
                    overflow = True
                    break
            if overflow:
                continue

        if not _replace_bullets_after(paragraphs, idx, bullets):
            logger.warning("Could not replace all 3 bullets for %r", name)

    # --- Skills swaps (replace low-signal skills with JD-relevant ones) ---
    _swap_skills(paragraphs, tailoring_data.get("skills_swaps", []))

    # --- Skills additions ---
    _append_skills(paragraphs, tailoring_data.get("skills_additions", []))

    # Save first pass
    doc.save(str(temp_docx))

    # --- One-page overflow check ---
    if _would_overflow_one_page(temp_docx):
        logger.warning("Tailored resume overflows one page — reverting to base resume with skills only")
        shutil.copy2(BASE_DOCX, temp_docx)
        # Re-apply only skills changes (they're small and won't cause overflow)
        doc = Document(str(temp_docx))
        _swap_skills(doc.paragraphs, tailoring_data.get("skills_swaps", []))
        _append_skills(doc.paragraphs, tailoring_data.get("skills_additions", []))
        doc.save(str(temp_docx))

    logger.info("Edited docx saved: %s", temp_docx.name)
    return temp_docx
